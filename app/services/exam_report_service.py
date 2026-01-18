# Exam Report Service - Professional DOCX/PDF reports with sub-questions and RTL support
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.units import inch, cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import io
from typing import Dict, Any, List, Optional
from datetime import datetime


class ExamReportService:
    """Generate professional exam reports in DOCX or PDF format."""
    
    LABELS = {
        'en': {
            'title': 'Exam Report',
            'student_name': 'Student Name',
            'section': 'Section',
            'date': 'Date',
            'subject': 'Subject',
            'final_score': 'Final Score',
            'question': 'Question',
            'sub_question': 'Part',
            'student_answer': 'Student Answer',
            'correct_answer': 'Correct Answer',
            'status': 'Status',
            'feedback': 'Feedback',
            'points': 'Points',
            'correct': 'Correct',
            'incorrect': 'Incorrect',
            'partial': 'Partial Credit'
        },
        'ar': {
            'title': 'تقرير الامتحان',
            'student_name': 'اسم الطالب',
            'section': 'الصف',
            'date': 'التاريخ',
            'subject': 'المادة',
            'final_score': 'الدرجة النهائية',
            'question': 'السؤال',
            'sub_question': 'الفرع',
            'student_answer': 'إجابة الطالب',
            'correct_answer': 'الإجابة الصحيحة',
            'status': 'الحالة',
            'feedback': 'ملاحظات',
            'points': 'الدرجات',
            'correct': 'صحيح',
            'incorrect': 'خطأ',
            'partial': 'جزئي'
        },
        'fr': {
            'title': 'Rapport d\'Examen',
            'student_name': 'Nom de l\'Élève',
            'section': 'Classe',
            'date': 'Date',
            'subject': 'Matière',
            'final_score': 'Note Finale',
            'question': 'Question',
            'sub_question': 'Partie',
            'student_answer': 'Réponse de l\'Élève',
            'correct_answer': 'Réponse Correcte',
            'status': 'Statut',
            'feedback': 'Commentaire',
            'points': 'Points',
            'correct': 'Correct',
            'incorrect': 'Incorrect',
            'partial': 'Crédit Partiel'
        }
    }
    
    def __init__(self):
        try:
            pdfmetrics.registerFont(TTFont('Arabic', 'C:/Windows/Fonts/arial.ttf'))
        except:
            pass
    
    def generate_report(self, student_info: Dict, grading_results: Dict,
                       questions: list, format: str = 'docx',
                       language: str = 'en') -> bytes:
        if format == 'pdf':
            return self._generate_pdf(student_info, grading_results, questions, language)
        return self._generate_docx(student_info, grading_results, questions, language)
    
    def _set_rtl(self, paragraph, is_rtl: bool):
        """Set RTL direction for Arabic."""
        if is_rtl:
            pPr = paragraph._p.get_or_add_pPr()
            bidi = OxmlElement('w:bidi')
            bidi.set(qn('w:val'), '1')
            pPr.append(bidi)
    
    def _generate_docx(self, student_info: Dict, grading_results: Dict,
                       questions: list, language: str) -> bytes:
        doc = Document()
        labels = self.LABELS.get(language, self.LABELS['en'])
        is_rtl = language == 'ar'
        
        # Get totals from request or calculate
        final_earned = grading_results.get('final_score', {}).get('earned',
                      grading_results.get('total_earned', 0))
        final_possible = grading_results.get('final_score', {}).get('possible',
                        grading_results.get('total_possible', 0))
        
        # === MODERN HEADER ===
        # Title
        title = doc.add_heading(labels['title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.color.rgb = RGBColor(30, 60, 120)
        self._set_rtl(title, is_rtl)
        
        # Subject & Date line
        subject = student_info.get('subject', '')
        date = student_info.get('date', datetime.now().strftime('%Y-%m-%d'))
        if subject:
            sub_para = doc.add_paragraph()
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_run = sub_para.add_run(f"{subject} • {date}")
            sub_run.font.size = Pt(11)
            sub_run.font.color.rgb = RGBColor(100, 100, 100)
            self._set_rtl(sub_para, is_rtl)
        
        # Final Score (smaller, elegant)
        score_para = doc.add_paragraph()
        score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        score_label = score_para.add_run(f"{labels['final_score']}: ")
        score_label.font.size = Pt(14)
        score_label.font.color.rgb = RGBColor(80, 80, 80)
        
        score_value = score_para.add_run(f"{final_earned}/{final_possible}")
        score_value.bold = True
        score_value.font.size = Pt(18)
        score_value.font.color.rgb = RGBColor(0, 120, 60) if final_earned >= final_possible * 0.6 else RGBColor(180, 50, 50)
        self._set_rtl(score_para, is_rtl)
        
        doc.add_paragraph()  # Spacer
        
        # === STUDENT INFO ===
        info_table = doc.add_table(rows=2, cols=4)
        info_table.style = 'Table Grid'
        
        headers = [labels['student_name'], labels['section'], labels['date'], labels['subject']]
        values = [
            student_info.get('name', 'N/A'),
            student_info.get('section', 'N/A'),
            student_info.get('date', datetime.now().strftime('%Y-%m-%d')),
            student_info.get('subject', 'N/A')
        ]
        
        for i, h in enumerate(headers):
            cell = info_table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                self._set_rtl(p, is_rtl)
        
        for i, v in enumerate(values):
            cell = info_table.rows[1].cells[i]
            cell.text = v
            for p in cell.paragraphs:
                self._set_rtl(p, is_rtl)
        
        doc.add_paragraph()
        
        # === QUESTIONS ===
        grading_list = grading_results.get('grading', [])
        
        for q_result in grading_list:
            q_num = q_result.get('question_number', '?')
            q_data = next((q for q in questions if str(q.get('question_number', '')) == str(q_num)), {})
            
            # Main question heading
            q_heading = doc.add_heading(f"{labels['question']} {q_num}", level=2)
            self._set_rtl(q_heading, is_rtl)
            
            q_text = q_data.get('question_text', q_data.get('statement', ''))
            if q_text:
                p = doc.add_paragraph(q_text)
                self._set_rtl(p, is_rtl)
            
            # Check for sub-questions
            sub_questions = q_result.get('sub_questions', q_result.get('branches', []))
            
            if sub_questions:
                # Process each sub-question
                for sub_idx, sub in enumerate(sub_questions):
                    sub_num = sub.get('sub_number', sub.get('branch', chr(97 + sub_idx)))  # a, b, c...
                    sub_heading = doc.add_heading(f"  {labels['sub_question']} ({sub_num})", level=3)
                    self._set_rtl(sub_heading, is_rtl)
                    
                    self._add_answer_table(doc, sub, labels, is_rtl)
            else:
                # Single question (no sub-questions)
                self._add_answer_table(doc, q_result, labels, is_rtl)
            
            doc.add_paragraph()
        
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()
    
    def _add_answer_table(self, doc, result: Dict, labels: Dict, is_rtl: bool):
        """Add answer details table."""
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        earned = float(result.get('earned_points', result.get('points_earned', 0)))
        possible = float(result.get('possible_points', result.get('points_possible', 1)))
        
        if earned >= possible and possible > 0:
            status = labels['correct']
        elif earned > 0:
            status = labels['partial']
        else:
            status = labels['incorrect']
        
        # For RTL languages, use LTR marks around numbers to fix bracket display
        if is_rtl:
            points_text = f"{status} \u200E({earned}/{possible})\u200E"  # LTR mark around brackets
        else:
            points_text = f"{status} ({earned}/{possible})"
        
        rows_data = [
            (labels['student_answer'], str(result.get('student_answer', 'N/A'))),
            (labels['correct_answer'], str(result.get('correct_answer', 'N/A'))),
            (labels['points'], points_text),
            (labels['feedback'], result.get('feedback', result.get('reason', '-')) or '-')
        ]
        
        for i, (label, value) in enumerate(rows_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
            for cell in table.rows[i].cells:
                for p in cell.paragraphs:
                    self._set_rtl(p, is_rtl)
                    if cell == table.rows[i].cells[0]:
                        for r in p.runs:
                            r.bold = True
    
    def _generate_pdf(self, student_info: Dict, grading_results: Dict,
                      questions: list, language: str) -> bytes:
        output = io.BytesIO()
        labels = self.LABELS.get(language, self.LABELS['en'])
        is_rtl = language == 'ar'
        
        doc = SimpleDocTemplate(output, pagesize=A4,
                               rightMargin=50, leftMargin=50,
                               topMargin=50, bottomMargin=50)
        
        styles = getSampleStyleSheet()
        story = []
        
        # Get totals
        final_earned = grading_results.get('final_score', {}).get('earned',
                      grading_results.get('total_earned', 0))
        final_possible = grading_results.get('final_score', {}).get('possible',
                        grading_results.get('total_possible', 0))
        
        # Title
        title_style = ParagraphStyle('Title', fontSize=24, alignment=1, spaceAfter=10)
        story.append(Paragraph(labels['title'], title_style))
        
        # Final Score (prominent)
        score_color = colors.darkgreen if final_earned >= final_possible * 0.6 else colors.darkred
        score_style = ParagraphStyle('Score', fontSize=28, alignment=1, 
                                     textColor=score_color, spaceAfter=20)
        story.append(Paragraph(f"{labels['final_score']}: {final_earned}/{final_possible}", score_style))
        story.append(Spacer(1, 15))
        
        # Student info
        info_data = [
            [labels['student_name'], student_info.get('name', 'N/A'),
             labels['subject'], student_info.get('subject', 'N/A')],
            [labels['section'], student_info.get('section', 'N/A'),
             labels['date'], student_info.get('date', datetime.now().strftime('%Y-%m-%d'))]
        ]
        info_table = Table(info_data, colWidths=[1.8*inch, 2*inch, 1.2*inch, 1.5*inch])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('BACKGROUND', (2, 0), (2, -1), colors.lightgrey),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTNAME', (2, 0), (2, -1), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        story.append(info_table)
        story.append(Spacer(1, 20))
        
        # Questions
        grading_list = grading_results.get('grading', [])
        heading_style = ParagraphStyle('Heading', fontSize=14, spaceBefore=15, spaceAfter=8)
        
        for q_result in grading_list:
            q_num = q_result.get('question_number', '?')
            q_data = next((q for q in questions if str(q.get('question_number', '')) == str(q_num)), {})
            
            story.append(Paragraph(f"<b>{labels['question']} {q_num}</b>", heading_style))
            
            q_text = q_data.get('question_text', '')
            if q_text:
                story.append(Paragraph(q_text, styles['Normal']))
                story.append(Spacer(1, 5))
            
            sub_questions = q_result.get('sub_questions', q_result.get('branches', []))
            
            if sub_questions:
                for sub_idx, sub in enumerate(sub_questions):
                    sub_num = sub.get('sub_number', chr(97 + sub_idx))
                    story.append(Paragraph(f"<i>{labels['sub_question']} ({sub_num})</i>", styles['Normal']))
                    self._add_pdf_answer_table(story, sub, labels)
            else:
                self._add_pdf_answer_table(story, q_result, labels)
            
            story.append(Spacer(1, 10))
        
        doc.build(story)
        output.seek(0)
        return output.getvalue()
    
    def _add_pdf_answer_table(self, story: list, result: Dict, labels: Dict):
        earned = float(result.get('earned_points', result.get('points_earned', 0)))
        possible = float(result.get('possible_points', result.get('points_possible', 1)))
        
        if earned >= possible and possible > 0:
            status = labels['correct']
        elif earned > 0:
            status = labels['partial']
        else:
            status = labels['incorrect']
        
        data = [
            [labels['student_answer'], str(result.get('student_answer', 'N/A'))],
            [labels['correct_answer'], str(result.get('correct_answer', 'N/A'))],
            [labels['points'], f"{status} ({earned}/{possible})"],
            [labels['feedback'], result.get('feedback', '-') or '-']
        ]
        
        table = Table(data, colWidths=[1.8*inch, 4.5*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        story.append(table)
        story.append(Spacer(1, 8))
